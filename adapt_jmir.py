from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from zipfile import ZipFile, ZIP_DEFLATED
import shutil

ROOT=Path(__file__).resolve().parent
src=ROOT/'CCSF_TobaccoControl_submission_ready.docx'
out=ROOT/'CCSF_JMIR_Infodemiology_Original_Paper.docx'
doc=Document(src)

def remove_para(p): p._element.getparent().remove(p._element)
def set_para(p,text,style=None):
    p.clear(); p.add_run(text)
    if style: p.style=style

def insert_after(p,text,style):
    new=doc.add_paragraph(text,style=style)
    p._p.addnext(new._p)
    return new

# Remove Tobacco Control-specific furniture and source abstract without relying on mutable paragraph indices.
paras=doc.paragraphs
start_idx=next(i for i,p in enumerate(paras) if p.text.strip()=="Abstract")
end_idx=next(i for i,p in enumerate(paras) if p.text.strip()=="What this paper adds")
for p in list(paras[start_idx+1:end_idx+1]): remove_para(p)
for table in list(doc.tables[:3]): table._element.getparent().remove(table._element)
# Clear journal headers/footers and keep only generic page furniture.
for section in doc.sections:
    for container in (section.header, section.footer):
        for p in container.paragraphs: p.clear()

# Front matter and structured abstract.
set_para(doc.paragraphs[0], 'An Account-Level Fingerprint for Prioritizing Potentially Coordinated Synthetic Advocacy in Alternative Nicotine Delivery Systems Discourse: Controlled Infodemiology Simulation Study')
set_para(doc.paragraphs[1], 'Original Paper')
abstract=next(p for p in doc.paragraphs if p.text.strip()=='Abstract')
prev=abstract
items=[
('Background','Commercial influence and apparently grassroots advocacy can shape online health-policy discourse about alternative nicotine delivery systems (ANDS). Infodemiology approaches could help prioritize account-level patterns for expert review, but transparent evidence is needed before any operational use.'),
('Objective','This study formulated Coordinated Communication Signal Fingerprinting (CCSF), an interpretable account-level framework for prioritizing potential patterns of coordinated synthetic advocacy in ANDS discourse, and tested its construct behavior in a controlled simulation.'),
('Methods','We analyzed a fully simulated, seeded corpus of 524 posts from 78 simulated accounts across 8 ANDS policy topics. The target generator was rule based rather than a modern large language model. Simulated organic users and simulated professional advocacy were comparison groups. Rule-perturbed synthetic posts were transformed programmatically; no human editing occurred. Account-level signals were language-model perplexity, sentence-length burstiness, commercial-policy framing density, and embedding convergence; stance variability was exploratory. We estimated within-corpus separation, ablation, a prevalence illustration, and exploratory clustering. A separate independent multi-generator validation framework was specified but has not produced research results.'),
('Results','The four-signal composite achieved AUC 0.995 (95% CI 0.982-1.000) for coordinated synthetic versus the 2 simulated human-style groups and AUC 0.955 for rule-perturbed synthetic versus simulated organic users. Perplexity alone achieved AUC 0.989; therefore, the small composite increment was not interpreted as evidence of superiority. Pipeline-based 5-fold random-fold interpolation was 0.980 (SD 0.030), and exploratory clustering was modest (adjusted Rand index 0.352). All results are within-corpus construct-validity evidence.'),
('Conclusions','CCSF offers a reproducible engineering baseline for infodemiology research on commercial influence in online health-policy discourse. It is not an externally validated detector or deployment-ready system. In any future operational setting, human expert review would be required before interpretation or action. Independent multi-generator and real-world validation remain necessary.')]
for heading,body in items:
    prev=insert_after(prev,heading,'Heading 2'); prev=insert_after(prev,body,'Normal')

# Locate heading/paragraphs by text and revise for JMIR framing.
def first(starts): return next(p for p in doc.paragraphs if p.text.startswith(starts))
set_para(first('Introduction'),'Introduction','Heading 1')
set_para(first('Methods'),'Methods','Heading 1')
set_para(first('Results'),'Results','Heading 1')
set_para(first('Discussion'),'Discussion','Heading 1')
set_para(first('RQ4:'),'What Does the Simulated Professional-Advocacy Control Reveal About Within-Corpus Specificity?','Heading 2')
set_para(first('Principal findings and their interpretation'),'Principal Findings','Heading 2')
set_para(first('Comparison with existing work'),'Comparison With Prior Work','Heading 2')
set_para(first('Strengths and limitations'),'Strengths and Limitations','Heading 2')
set_para(first('Implications and future work'),'Implications for Infodemiology Research','Heading 2')
conclusion=first('Conclusion')
set_para(conclusion,'Conclusions','Heading 2')

# Definition and exact JMIR language in key text/legends.
repls={
'We propose CCSF,':'We propose Coordinated Communication Signal Fingerprinting (CCSF),',
'organic users':'simulated organic users',
'Professional advocacy':'Simulated professional advocacy',
'professional advocacy':'simulated professional advocacy',
'Organic users':'Simulated organic users',
'coordinated synthetic accounts':'coordinated synthetic accounts',
'Figure 2. Distributions of the four primary account-level features by constructed group.':'Figure 2. Distributions of the four primary account-level features by group: Simulated Organic, Coordinated Synthetic, Simulated Professional, and Rule-Perturbed Synthetic.',
'Figure 3. Standardised baseline fingerprint profiles by constructed group.':'Figure 3. Standardized baseline fingerprint profiles for Simulated Organic, Coordinated Synthetic, Simulated Professional, and Rule-Perturbed Synthetic groups.',
}
for p in doc.paragraphs:
    text=p.text; new=text
    for a,b in repls.items(): new=new.replace(a,b)
    if new!=text: set_para(p,new)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text=p.text; new=text
                for a,b in repls.items(): new=new.replace(a,b)
                if new!=text: set_para(p,new)

# Update table labels after the first three journal-specific boxes were removed.
t1,t2,t3=doc.tables
for row in t1.rows:
    if row.cells[0].text=='Simulated organic users': row.cells[0].text='Simulated Organic'
    if row.cells[0].text=='Coordinated synthetic': row.cells[0].text='Coordinated Synthetic'
    if 'professional' in row.cells[0].text.lower(): row.cells[0].text='Simulated Professional'
    if 'Rule-perturbed' in row.cells[0].text: row.cells[0].text='Rule-Perturbed Synthetic'
t2.cell(0,1).text='Simulated Organic (mean±SD)'; t2.cell(0,2).text='Coordinated Synthetic (mean±SD)'; t2.cell(0,3).text='Simulated Professional (mean±SD)'

# Replace old declarations with JMIR-compatible end matter before References.
refs=first('References')
# remove declaration heading and its body until references
start=next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='Declarations')
end=next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='References')
for p in list(doc.paragraphs[start:end]): remove_para(p)
prev=refs._p.getprevious()
# Resolve paragraph wrapper after deletion
anchor=next(p for p in doc.paragraphs if p.text.strip()=='References')
endmatter=[
('Acknowledgments','None.'),
('Funding','This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.'),
('Conflicts of Interest','None declared.'),
('Data Availability','The synthetic baseline corpus, generator, analysis code, derived outputs, and verification materials are included with this manuscript for review. A public repository URL, tagged release, persistent archive, DOI, and license identifiers must be inserted before submission once the release exists; no DOI or public archive is claimed in this version.'),
('Authors’ Contributions','Conceptualization: SE. Data Curation: SE. Formal Analysis: SE. Investigation: SE. Methodology: SE. Project Administration: SE. Software: SE. Validation: SE. Visualization: SE. Writing – Original Draft: SE. Writing – Review & Editing: SE.'),
('Abbreviations','ANDS: alternative nicotine delivery systems; AUC: area under the receiver operating characteristic curve; CCSF: Coordinated Communication Signal Fingerprinting; PPV: positive predictive value.'),
]
prev=anchor
for heading,body in reversed(endmatter):
    p=doc.add_paragraph(body,style='Normal'); anchor._p.addprevious(p._p)
    h=doc.add_paragraph(heading,style='Heading 1'); p._p.addprevious(h._p)

# Correct future-operation grammar and final group nomenclature.
for p in doc.paragraphs:
    new=p.text.replace('not a automatic','not an automatic').replace('not a an automatic','not an automatic')
    if new!=p.text: set_para(p,new)
# Add accessible headers/alt metadata.
for table in doc.tables:
    trPr=table.rows[0]._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader'); hdr.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val','true'); trPr.append(hdr)
alt=['Account-level CCSF workflow.','Four account-level feature distributions by simulated group.','Standardized fingerprint profiles by simulated group.','Illustrative t-SNE projection of post embeddings.','Within-corpus ROC and ablation results.','Lexical overlap and semantic similarity by simulated group.']
for shape,description in zip(doc.inline_shapes,alt): shape._inline.docPr.set('descr',description); shape._inline.docPr.set('title',description)

doc.save(out)
# replace figures in package
figs=['fig1_pipeline.png','fig2_features.png','fig3_fingerprint.png','fig4_tsne.png','fig5_roc_ablation.png','fig6_nonduplication.png']
with ZipFile(out,'r') as zin:
    temp=out.with_suffix('.tmp.docx')
    with ZipFile(temp,'w',ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            n=item.filename
            if n.startswith('word/media/image') and n.endswith('.png'):
                num=int(n.rsplit('image',1)[1].split('.',1)[0])
                if 1<=num<=6: zout.writestr(item,(ROOT/figs[num-1]).read_bytes()); continue
            zout.writestr(item,zin.read(n))
shutil.move(temp,out)
print(out)
